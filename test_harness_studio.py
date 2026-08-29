# -*- coding: utf-8 -*-
"""Nova Harness, profilo «studio»: aprire, trovare, leggere intorno.

La cosa che si prova qui non e' che sappia rispondere: e' che sappia
**indicare**. Una ricerca che torna una posizione o e' giusta o e' sbagliata,
e si vede; una che torna una frase puo' essere inventata benissimo. E' per
questo che il profilo esiste.

Si costruisce un .docx e un .pdf veri, sul momento: provare l'ancoraggio su
un finto documento sarebbe provare niente.
"""
import os
import sys
import tempfile
from pathlib import Path

RADICE = Path(__file__).resolve().parent
sys.path.insert(0, str(RADICE))
sys.stdout.reconfigure(encoding="utf-8", errors="replace")

finto = Path(tempfile.mkdtemp(prefix="nova_harness_"))
os.environ["APPDATA"] = str(finto)

from nova import harness           # noqa: E402
from nova import harness_finestra  # noqa: E402

# Questa prova riguarda l'ancoraggio, non il vetro: la finestra non si apre.
harness_finestra.apri_se_serve = lambda *a, **k: {
    "viva": False, "accesa_adesso": False, "motivo": "non in una prova"}

passati = 0
falliti: list[str] = []


def controlla(nome, condizione, dettaglio=""):
    global passati
    if condizione:
        passati += 1
        print(f"  [ok ] {nome}")
    else:
        falliti.append(nome)
        print(f"  [NO ] {nome}  {dettaglio}")


TESTI = [
    ("Titolo", "Stato dell'arte dei modelli open source"),
    ("Normale", "I modelli a pesi aperti hanno colmato buona parte del "
                "divario con quelli chiusi nel corso del 2025."),
    ("Normale", "La quantizzazione a quattro bit riduce la memoria "
                "necessaria di circa il settanta per cento, con una perdita "
                "di qualita' che resta contenuta."),
    ("Normale", "Engram introduce una memoria condizionale indicizzata da "
                "n-grammi, letta in tempo costante dentro il passaggio in "
                "avanti del modello."),
    ("Normale", "Le licenze restano il punto piu' controverso: aperto non "
                "vuol dire libero, e diverse famiglie impongono limiti d'uso."),
]

# --- un .docx vero ----------------------------------------------------
import docx  # noqa: E402
d = docx.Document()
for stile, t in TESTI:
    d.add_paragraph(t, style="Title" if stile == "Titolo" else None)
tab = d.add_table(rows=1, cols=2)
tab.rows[0].cells[0].text = "Llama"
tab.rows[0].cells[1].text = "licenza comunitaria"
doc_word = finto / "ricerca.docx"
d.save(str(doc_word))

# --- un .pdf vero -----------------------------------------------------
import fitz  # noqa: E402
p = fitz.open()
pagina = p.new_page()
y = 72
for _, t in TESTI:
    pagina.insert_textbox(fitz.Rect(56, y, 540, y + 60), t, fontsize=11)
    y += 70
seconda = p.new_page()
seconda.insert_textbox(fitz.Rect(56, 72, 540, 140),
                       "La firma digitale qualificata ha valore legale "
                       "equiparato alla firma autografa.", fontsize=11)
doc_pdf = finto / "ricerca.pdf"
p.save(str(doc_pdf))
p.close()


print("\n1. si rifiuta di aprire quello che non sa aprire")
r = harness.apri(str(finto / "non-esiste.docx"))
controlla("un file che non c'e'", not r.get("ok") and "non trovo" in r["motivo"])
(finto / "roba.xyz").write_text("boh", encoding="utf-8")
r = harness.apri(str(finto / "roba.xyz"))
controlla("un formato sconosciuto, dicendo quali sa fare",
          not r.get("ok") and "So aprire" in r["motivo"], str(r))
r = harness.apri(str(doc_word), profilo="inventato")
controlla("un profilo che non esiste", not r.get("ok") and "sconosciuto" in r["motivo"])
controlla("e cercare senza aver aperto niente non finge",
          not harness.cerca("qualcosa").get("ok"))

print("\n2. il documento Word, fatto a pezzi")
r = harness.apri(str(doc_word))
controlla("si apre", r.get("ok"), str(r.get("motivo")))
controlla("con un blocco per paragrafo, tabella compresa",
          r.get("blocchi") == len(TESTI) + 1, str(r.get("blocchi")))
s = harness.stato()
controlla("e lo stato sa qual e'", s.get("nome") == "ricerca.docx", str(s))

print("\n3. cercare vuol dire indicare, non raccontare")
r = harness.cerca("quanto risparmia la quantizzazione")
controlla("trova qualcosa", r.get("ok") and r.get("trovati"), str(r)[:120])
primo = r["trovati"][0]
controlla("e il blocco indicato contiene davvero la risposta",
          "settanta per cento" in primo["testo"], primo["testo"][:70])
controlla("con un identificativo utilizzabile", primo["id"].startswith("p"))
controlla("e un punteggio", 0 < primo["quanto"] <= 1)

r = harness.cerca("engram n-grammi memoria")
controlla("trova anche il paragrafo su Engram",
          "Engram" in r["trovati"][0]["testo"], r["trovati"][0]["testo"][:60])

r = harness.cerca("quantizzaziome")     # scritto storto di proposito
controlla("e sopravvive a un refuso, come le ricette",
          r.get("trovati") and "settanta" in r["trovati"][0]["testo"],
          str([t["testo"][:40] for t in (r.get("trovati") or [])[:2]]))

r = harness.cerca("ricetta della carbonara")
controlla("quello che non c'e' non lo inventa",
          not r.get("trovati"), str(r.get("trovati")))

print("\n4. quello che si trova viene evidenziato")
harness.cerca("licenze aperto libero")
s = harness.stato()
controlla("lo stato porta gli evidenziati", bool(s.get("evidenziati")), str(s))

print("\n5. leggere intorno a un punto")
r = harness.cerca("engram")
pid = r["trovati"][0]["id"]
c = harness.leggi(intorno=pid, blocchi=1)
controlla("torna il contesto attorno", c.get("ok") and "Engram" in c["testo"])
controlla("e anche cio' che gli sta accanto",
          "quantizzazione" in c["testo"].lower() or "licenze" in c["testo"].lower(),
          c["testo"][:100])
controlla("un punto inesistente viene detto, non ignorato",
          not harness.leggi(intorno="p999").get("ok"))

print("\n6. il PDF: qui la posizione ha un riquadro")
r = harness.apri(str(doc_pdf))
controlla("si apre", r.get("ok"), str(r.get("motivo")))
controlla("e sa quante pagine sono", r.get("pagine") == 2, str(r.get("pagine")))
r = harness.cerca("firma digitale valore legale")
t = r["trovati"][0]
controlla("trova nella pagina giusta", t["pagina"] == 2, str(t["pagina"]))
s = harness._stato()
blocco = [b for b in s["blocchi"] if b["id"] == t["id"]][0]
controlla("e il blocco porta il riquadro per evidenziarlo",
          blocco["riquadro"] and len(blocco["riquadro"]) == 4,
          str(blocco["riquadro"]))

print("\n6b. il PDF si vede com'e', non come testo estratto")
pagine = harness.pagine_disegnate()
controlla("le pagine vengono disegnate", len(pagine) == 2, str(len(pagine)))
controlla("i file esistono davvero e non sono vuoti",
          all(Path(p["file"]).is_file() and Path(p["file"]).stat().st_size > 1000
              for p in pagine),
          str([(Path(p["file"]).name, Path(p["file"]).stat().st_size)
               for p in pagine if Path(p["file"]).exists()]))
# La ricerca di prima aveva trovato la firma digitale, che sta a pagina 2.
accese = [p["pagina"] for p in pagine if p["evidenziata"]]
controlla("e l'evidenziazione sta sulla pagina giusta", accese == [2], str(accese))

prima = {p["pagina"]: Path(p["file"]).name for p in pagine}
harness.cerca("quantizzazione quattro bit")
dopo = {p["pagina"]: Path(p["file"]).name for p in harness.pagine_disegnate()}
controlla("cambiando ricerca cambia la pagina evidenziata",
          dopo[1] != prima[1] or dopo[2] != prima[2], f"{prima} -> {dopo}")
controlla("e le pagine si tengono da parte invece di rifarle ogni volta",
          len(list((harness._base() / f"{harness._stato()['sessione']}-pagine")
                   .glob("*.png"))) > 2)

controlla("per un documento senza pagine non si finge una pagina",
          harness.apri(str(doc_word)).get("ok") and harness.pagine_disegnate() == [])
# Aprire il Word ha aperto una sessione nuova: si rimette il PDF, o la
# sezione dopo leggerebbe un registro appena nato e vuoto.
harness.apri(str(doc_pdf))
harness.cerca("firma digitale valore legale")

print("\n7. la sessione resta scritta, non solo a schermo")
ev = harness.eventi()
tipi = [x["evento"] for x in ev]
controlla("il registro c'e'", bool(ev))
controlla("e racconta cosa e' stato aperto e cercato",
          "aperto" in tipi and "cercato" in tipi, str(tipi))
controlla("con la domanda vera dentro",
          any("firma" in (x.get("domanda") or "") for x in ev), str(ev[-2:]))
controlla("chiudere si puo'", harness.chiudi())
controlla("e dopo non c'e' piu' niente di aperto",
          not harness.stato().get("ok"))

print("\nX. una pila di documenti si cerca tutta insieme")
# La differenza fra leggere un libro e studiare su sei: la domanda vera non
# e' «dove sta in questo file» ma «in quale file sta».
import tempfile as _tf
biblioteca = Path(_tf.mkdtemp(prefix="nova_bib_"))
(biblioteca / "libro_uno.md").write_text(
    "# Termodinamica\n\nIl primo principio dice che l'energia si conserva.\n\n"
    "L'entropia non compare qui.\n", encoding="utf-8")
(biblioteca / "libro_due.md").write_text(
    "# Statistica\n\nLa distribuzione normale ha due parametri.\n\n"
    "Il teorema del limite centrale non riguarda l'entropia.\n",
    encoding="utf-8")
(biblioteca / "appunti.md").write_text(
    "# Appunti\n\nL'entropia di Boltzmann lega il macrostato ai microstati.\n",
    encoding="utf-8")

r = harness.apri(str(biblioteca))
controlla("la cartella si apre come progetto", r.get("ok"), str(r.get("motivo")))
controlla("con dentro tre testi", r.get("file_nel_progetto") == 3,
          str(r.get("file_nel_progetto")))

d = harness.cerca_progetto("entropia di Boltzmann microstati")
controlla("la ricerca sul progetto risponde", d.get("ok"), str(d.get("motivo")))
controlla("ha guardato in tutti e tre", d.get("cercati") == 3, str(d.get("cercati")))
controlla("e il primo posto e' negli appunti",
          d["risultati"][0]["file"] == "appunti.md",
          str([x["file"] for x in d["risultati"]]))
controlla("dice il blocco, non solo il file",
          bool(d["risultati"][0]["blocco"]), str(d["risultati"][0]))

# Cercare qualcosa che non c'e' deve dirlo, non consegnare il meno peggio.
vuoto = harness.cerca_progetto("kubernetes ingress controller")
controlla("quello che non c'e' non si trova",
          vuoto.get("ok") and not vuoto["risultati"], str(vuoto))

# L'indice non deve rileggere quello che non e' cambiato, ma deve accorgersi
# di quello che cambia: altrimenti si studia su una copia vecchia.
(biblioteca / "libro_uno.md").write_text(
    "# Termodinamica\n\nIl secondo principio parla proprio di entropia "
    "crescente.\n", encoding="utf-8")
d2 = harness.cerca_progetto("entropia crescente secondo principio")
controlla("una modifica al testo entra subito nell'indice",
          any(x["file"] == "libro_uno.md" for x in d2["risultati"]),
          str([x["file"] for x in d2["risultati"]]))

fuori = harness.cerca_progetto("")
controlla("una domanda vuota si rifiuta", fuori.get("ok") is False)

harness.apri(str(biblioteca / "libro_due.md"))
solo = harness.cerca_progetto("entropia")
controlla("passando a un file, il progetto resta cercabile",
          solo.get("ok") and solo.get("cercati") == 3, str(solo.get("motivo")))

print(f"\n{passati}/{passati + len(falliti)} passati")
for x in falliti:
    print("  FALLITO:", x)
sys.exit(1 if falliti else 0)
