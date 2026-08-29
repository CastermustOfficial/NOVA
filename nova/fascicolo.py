# -*- coding: utf-8 -*-
"""Il fascicolo: i fatti veri sull'utente, in un posto solo.

Perche' esiste. NOVA scrive candidature al posto dell'utente. Ogni volta
deve ricostruire chi e' - dove ha lavorato, cosa sa fare, come scrive - e un
modello a cui manca un dato non lo lascia vuoto: lo riempie. Un refuso in una
lettera di presentazione e' un refuso; un'esperienza inventata e' una
dichiarazione falsa a un datore di lavoro, con sopra la firma dell'utente.

Il rimedio non e' una regola in piu' sul non inventare - quelle si obbediscono
finche' c'e' il dato. E' avere il dato: una cartella dove stanno il CV, le
esperienze vere, i testi che l'utente ha gia' scritto di suo pugno. Da li' si
pesca, e quello che li' non c'e' **si chiede**, non si deduce.

E' anche il posto giusto per la voce: chi ha scritto tre lettere di
presentazione ne ha gia' il tono. Ricopiarlo e' meglio che immaginarlo.

Sta in `Documenti\\NOVA\\fascicolo` e non sotto %APPDATA%, di proposito: e'
roba dell'utente, deve trovarla, aprirla e cambiarla senza sapere dove NOVA
tiene le proprie cose.
"""
from __future__ import annotations

import json
import os
from datetime import datetime
from pathlib import Path

# Quello che si sa leggere. Il resto si elenca lo stesso - sapere che un file
# c'e' vale anche quando non lo si sa aprire - ma dicendolo.
TESTO = {".txt", ".md", ".markdown", ".json", ".csv", ".tsv", ".yaml", ".yml",
         ".html", ".htm", ".log", ".rst"}
LEGGIBILI = TESTO | {".pdf", ".docx", ".xlsx"}


def cartella() -> Path:
    """Dove sta il fascicolo. Si puo' spostare da `fascicolo` in config.json."""
    try:
        from .config import Config
        scelta = (Config.load().fascicolo or "").strip()
        if scelta:
            return Path(scelta).expanduser()
    except Exception:
        pass
    documenti = Path.home() / "Documents"
    if not documenti.is_dir():
        alt = Path.home() / "Documenti"
        documenti = alt if alt.is_dir() else Path.home()
    return documenti / "NOVA" / "fascicolo"


def prepara() -> Path:
    """Crea la cartella se non c'e', con dentro una riga che spiega a cosa serve."""
    c = cartella()
    c.mkdir(parents=True, exist_ok=True)
    guida = c / "LEGGIMI.md"
    if not guida.exists():
        guida.write_text(
            "# Il fascicolo\n\n"
            "Qui dentro vanno i fatti veri su di te: il CV, le esperienze, i\n"
            "progetti, i testi che hai gia' scritto tu.\n\n"
            "NOVA pesca da qui quando scrive qualcosa a nome tuo - una\n"
            "candidatura, una lettera, una biografia. Quello che qui non c'e'\n"
            "**te lo chiede**, invece di dedurlo: e' la differenza fra un\n"
            "errore e una dichiarazione falsa con sopra la tua firma.\n\n"
            "Legge .txt, .md, .pdf, .docx, .xlsx, .csv e .json. Puoi\n"
            "organizzarlo in sottocartelle come preferisci.\n",
            encoding="utf-8")
    return c


def elenco(massimo: int = 100) -> list[dict]:
    """Cosa c'e' nel fascicolo, con dimensione e data."""
    c = cartella()
    if not c.is_dir():
        return []
    fuori: list[dict] = []
    for f in sorted(c.rglob("*")):
        if not f.is_file() or f.name.startswith("~$"):
            continue
        try:
            st = f.stat()
        except OSError:
            continue
        fuori.append({
            "nome": str(f.relative_to(c)).replace("\\", "/"),
            "byte": st.st_size,
            "quando": datetime.fromtimestamp(st.st_mtime).strftime("%d/%m/%Y"),
            "leggibile": f.suffix.lower() in LEGGIBILI,
        })
        if len(fuori) >= massimo:
            break
    return fuori


def _da_pdf(f: Path, pagine: int = 40) -> str:
    from pypdf import PdfReader
    r = PdfReader(str(f))
    pezzi = [(p.extract_text() or "") for p in r.pages[:pagine]]
    return "\n\n".join(x.strip() for x in pezzi if x.strip())


def _da_docx(f: Path) -> str:
    import docx
    d = docx.Document(str(f))
    righe = [p.text for p in d.paragraphs]
    for t in d.tables:
        for r in t.rows:
            righe.append("\t".join(c.text.strip() for c in r.cells))
    return "\n".join(x for x in righe if x.strip())


def _da_xlsx(f: Path, righe_max: int = 300) -> str:
    import openpyxl
    w = openpyxl.load_workbook(str(f), data_only=True, read_only=True)
    fuori = []
    for foglio in w.worksheets:
        fuori.append(f"# {foglio.title}")
        for i, riga in enumerate(foglio.iter_rows(values_only=True)):
            if i >= righe_max:
                fuori.append("[...]")
                break
            fuori.append("\t".join("" if v is None else str(v) for v in riga))
    return "\n".join(fuori)


def leggi(nome: str, caratteri: int = 8000) -> dict:
    """Il contenuto di un pezzo del fascicolo, come testo."""
    c = cartella()
    f = (c / nome).resolve()
    try:
        # Un nome con «..» dentro non deve poter uscire dalla cartella: il
        # fascicolo e' un permesso di leggere quella, non il disco.
        f.relative_to(c.resolve())
    except ValueError:
        return {"ok": False, "motivo": "quel nome esce dal fascicolo"}
    if not f.is_file():
        return {"ok": False, "motivo": f"«{nome}» non c'e' nel fascicolo"}
    est = f.suffix.lower()
    try:
        if est in TESTO:
            testo = f.read_text(encoding="utf-8", errors="replace")
        elif est == ".pdf":
            testo = _da_pdf(f)
        elif est == ".docx":
            testo = _da_docx(f)
        elif est == ".xlsx":
            testo = _da_xlsx(f)
        else:
            return {"ok": False,
                    "motivo": f"non so leggere un {est or 'file senza estensione'}. "
                              "Si legge .txt .md .pdf .docx .xlsx .csv .json"}
    except Exception as e:
        return {"ok": False, "motivo": f"{type(e).__name__}: {e}"}
    if not (testo or "").strip():
        return {"ok": False,
                "motivo": "il file non contiene testo estraibile "
                          "(se e' un PDF scansionato, servirebbe un OCR)"}
    return {"ok": True, "nome": nome, "caratteri": len(testo),
            "testo": testo[:caratteri], "tagliato": len(testo) > caratteri}


def indice(massimo: int = 100) -> str:
    """L'elenco in una forma che si legge."""
    voci = elenco(massimo)
    if not voci:
        return (f"Il fascicolo e' vuoto ({cartella()}).\n"
                "Finche' e' vuoto, di fatti sull'utente non ne hai: chiediglieli "
                "invece di dedurli.")
    righe = [f"{len(voci)} file in {cartella()}:"]
    for v in voci:
        kb = v["byte"] / 1024
        nota = "" if v["leggibile"] else "   (non so leggerlo)"
        righe.append(f"  {v['nome']}  —  {kb:.0f} KB, {v['quando']}{nota}")
    return "\n".join(righe)
