"""Leggere documenti che non sono testo semplice.

NOVA sapeva aprire un .txt e non una fattura in PDF: e' la lacuna che si
incontra prima di tutte, perche' i documenti che contano sul computer di
qualcuno quasi mai sono file di testo.

Nessuna dipendenza nuova: pypdf, PyMuPDF, python-docx e openpyxl sono gia'
nell'ambiente. Se una manca, si dice quale e come installarla invece di
fallire con un errore di importazione.
"""
from __future__ import annotations

from pathlib import Path

from .base import Risk, ToolError, tool

# Oltre questa soglia si taglia: un PDF di trecento pagine riempirebbe il
# contesto e lascerebbe il modello senza spazio per ragionarci sopra.
CARATTERI_MASSIMI = 30_000


def _taglia(testo: str, quante_pagine: int | None = None) -> str:
    if len(testo) <= CARATTERI_MASSIMI:
        return testo
    tagliato = testo[:CARATTERI_MASSIMI]
    coda = f"\n\n[...documento troncato a {CARATTERI_MASSIMI} caratteri"
    if quante_pagine:
        coda += f" su {quante_pagine} pagine"
    coda += ". Chiedi una parte precisa con «pagine» per vedere il resto.]"
    return tagliato + coda


def _pdf(p: Path, pagine: str) -> str:
    try:
        import pypdf
    except ImportError:
        raise ToolError("per i PDF serve pypdf: pip install pypdf")
    lettore = pypdf.PdfReader(str(p))
    if lettore.is_encrypted:
        try:
            lettore.decrypt("")
        except Exception:
            raise ToolError(f"{p.name} e' protetto da password: non riesco ad aprirlo")
    totale = len(lettore.pages)
    indici = _intervallo(pagine, totale)
    pezzi = []
    for i in indici:
        try:
            t = lettore.pages[i].extract_text() or ""
        except Exception as e:
            t = f"[pagina {i+1} illeggibile: {e}]"
        pezzi.append(f"--- pagina {i + 1} di {totale} ---\n{t.strip()}")
    testo = "\n\n".join(pezzi).strip()
    if not testo or len(testo.replace("---", "").strip()) < 20:
        raise ToolError(
            f"{p.name} non contiene testo estraibile: e' probabilmente una scansione. "
            "Serve il riconoscimento ottico, che NOVA non ha ancora."
        )
    return _taglia(testo, totale)


def _intervallo(pagine: str, totale: int) -> list[int]:
    """«3», «2-5», «» -> tutte."""
    pagine = (pagine or "").strip()
    if not pagine:
        return list(range(totale))
    try:
        if "-" in pagine:
            a, b = pagine.split("-", 1)
            inizio, fine = int(a) - 1, int(b)
        else:
            inizio, fine = int(pagine) - 1, int(pagine)
    except ValueError:
        raise ToolError("«pagine» vuole un numero (3) o un intervallo (2-5)")
    inizio = max(0, inizio)
    fine = min(totale, fine)
    if inizio >= fine:
        raise ToolError(f"il documento ha {totale} pagine: «{pagine}» non ci sta dentro")
    return list(range(inizio, fine))


def _docx(p: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError:
        raise ToolError("per i .docx serve python-docx: pip install python-docx")
    d = docx.Document(str(p))
    pezzi = [par.text for par in d.paragraphs if par.text.strip()]
    # Le tabelle sono spesso il contenuto vero di un documento di lavoro:
    # ignorarle vorrebbe dire leggere una fattura senza gli importi.
    for n, tab in enumerate(d.tables, 1):
        righe = []
        for r in tab.rows:
            celle = [c.text.strip() for c in r.cells]
            if any(celle):
                righe.append(" | ".join(celle))
        if righe:
            pezzi.append(f"\n--- tabella {n} ---\n" + "\n".join(righe))
    return _taglia("\n".join(pezzi).strip())


def _xlsx(p: Path, foglio: str) -> str:
    try:
        import openpyxl
    except ImportError:
        raise ToolError("per i fogli di calcolo serve openpyxl: pip install openpyxl")
    # data_only: interessa il risultato, non la formula che lo produce.
    wb = openpyxl.load_workbook(str(p), data_only=True, read_only=True)
    nomi = wb.sheetnames
    da_leggere = [foglio] if foglio else nomi
    if foglio and foglio not in nomi:
        raise ToolError(f"in questo file non c'e' un foglio «{foglio}». Ci sono: {', '.join(nomi)}")
    pezzi = []
    for nome in da_leggere:
        ws = wb[nome]
        righe = []
        for r in ws.iter_rows(values_only=True):
            celle = ["" if c is None else str(c) for c in r]
            if any(c.strip() for c in celle):
                righe.append(" | ".join(celle))
            if len(righe) > 500:
                righe.append("[...foglio troncato a 500 righe]")
                break
        if righe:
            pezzi.append(f"--- foglio «{nome}» ---\n" + "\n".join(righe))
    wb.close()
    if not pezzi:
        raise ToolError(f"{p.name} e' vuoto")
    return _taglia("\n\n".join(pezzi))


@tool(
    "read_document",
    "Legge il CONTENUTO di un PDF, Word (.docx), Excel (.xlsx) o file di testo. "
    "Usa questo invece di read_file quando il documento non e' testo semplice: "
    "read_file su un PDF restituisce byte illeggibili.",
    {
        "path": {"type": "string", "description": "Percorso del documento"},
        "pagine": {"type": "string",
                   "description": "Solo per i PDF: «3» o «2-5». Vuoto = tutto"},
        "foglio": {"type": "string",
                   "description": "Solo per i fogli di calcolo: nome del foglio. Vuoto = tutti"},
    },
    Risk.SAFE, required=["path"], category="file",
    preview=lambda a: f"Legge il contenuto di {a.get('path')}",
)
def read_document(path: str, pagine: str = "", foglio: str = "") -> str:
    p = Path(path).expanduser()
    if not p.exists():
        raise ToolError(f"{p} non esiste")
    if p.is_dir():
        raise ToolError(f"{p} e' una cartella, non un documento")
    est = p.suffix.lower()
    if est == ".pdf":
        return _pdf(p, pagine)
    if est in (".docx", ".docm"):
        return _docx(p)
    if est in (".xlsx", ".xlsm"):
        return _xlsx(p, foglio)
    if est == ".doc":
        raise ToolError(
            ".doc e' il vecchio formato di Word e non si legge senza Word installato. "
            "Aprilo e salvalo come .docx, oppure dimmi di provare con Word."
        )
    # Tutto il resto si tenta come testo: e' meglio provare che rifiutare per
    # via dell'estensione, perche' meta' dei file di configurazione non ne ha
    # una riconoscibile.
    try:
        return _taglia(p.read_text(encoding="utf-8", errors="replace"))
    except Exception as e:
        raise ToolError(f"non so leggere {p.name} ({est or 'senza estensione'}): {e}")