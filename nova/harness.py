# -*- coding: utf-8 -*-
"""Nova Harness — il posto dove il lavoro si deposita.

La chat e' dove si parla; l'harness e' dove il lavoro sta. Alla chat il
verdetto - due righe: cosa ho fatto, cosa ho trovato, cosa devi decidere -
all'harness il materiale.

Questo file e' la meta' che non si vede: la sessione, il documento fatto a
pezzi, e la ricerca che torna una POSIZIONE invece di un'affermazione.

Perche' la posizione e' la funzione, e non un vezzo dell'interfaccia. Un
assistente allo studio che dice «lo trovi a pagina 12, terzo blocco» non puo'
bluffare: o quel blocco contiene quella cosa o non la contiene. E' la stessa
medicina del fascicolo - i fatti vengono da un posto reale - applicata alla
lettura invece che alla scrittura.

La sessione e' un registro **in aggiunta**, su disco. Serve a due cose: che
la chat possa dire «l'ho messo li'» e che «li'» esista ancora domani; e che
si possa rileggere cosa e' stato guardato, senza fidarsi del racconto. E'
l'invariante di deepseek-harness - cio' che il modello ha visto dev'essere
ricostruibile - applicata dove ci serviva davvero.
"""
from __future__ import annotations

import hashlib
import json
import os
import re
import time
import unicodedata
import uuid
from datetime import datetime
from pathlib import Path

# I profili: quali strumenti e quanta istruzione servono per un mestiere.
# Non e' eleganza architetturale, e' la ragione per cui un modello piccolo
# puo' fare questo lavoro: «studio» chiede cinque strumenti, non ventitre.
PROFILI = {
    "studio": {
        "titolo": "Studio",
        "descrizione": "Un documento a sinistra, la chat a destra. Si legge, "
                       "si cerca, si evidenzia. Non si scrive.",
        "strumenti": ["harness_apri", "harness_cerca", "harness_leggi",
                      "harness_stato"],
        "scrive": False,
    },
}

CODICE = {".py", ".js", ".mjs", ".ts", ".jsx", ".tsx", ".css", ".json",
          ".yml", ".yaml", ".toml", ".ini", ".cfg", ".rs", ".go", ".java",
          ".c", ".h", ".cpp", ".hpp", ".cs", ".rb", ".php", ".sh", ".bat",
          ".ps1", ".sql", ".xml", ".svg", ".vue", ".svelte", ".gitignore"}
A_RIGHE = CODICE | {".html", ".htm"}
LEGGIBILI = {".pdf", ".docx", ".txt", ".md", ".html", ".htm"} | CODICE

# Cartelle che in un progetto non si guardano: sono il prodotto, non il
# lavoro, e riempirebbero l'albero di roba che nessuno apre.
NON_GUARDARE = {".git", "node_modules", "__pycache__", ".venv", "venv",
                "dist", "build", ".mypy_cache", ".pytest_cache", ".idea",
                ".vscode", "target", ".next", "site-packages"}
FILE_MAX = 400_000
ALBERO_MAX = 600
# Un artifact e' una pagina, e una pagina si guarda resa, non
# spogliata: la finestra la disegna davvero. I blocchi pero' restano
# le righe del sorgente, perche' quello che NOVA deve poter cambiare
# e' il codice, non il testo che si vede.


def _base() -> Path:
    b = os.environ.get("APPDATA")
    return (Path(b) / "NOVA" if b else Path.home() / ".config" / "NOVA") / "harness"


def _puntatore() -> Path:
    """Dove sta scritto qual e' la sessione aperta adesso: lo legge la finestra."""
    return _base() / "corrente.json"


def _registro(sessione: str) -> Path:
    return _base() / f"{sessione}.jsonl"


def _annota(sessione: str, evento: str, **dati) -> None:
    try:
        f = _registro(sessione)
        f.parent.mkdir(parents=True, exist_ok=True)
        with open(f, "a", encoding="utf-8") as fh:
            fh.write(json.dumps(
                {"quando": datetime.now().isoformat(timespec="seconds"),
                 "evento": evento, **dati}, ensure_ascii=False) + "\n")
    except Exception:
        pass


# ------------------------------------------------------ il documento a pezzi

def _blocchi_docx(f: Path) -> list[dict]:
    import docx
    d = docx.Document(str(f))
    fuori = []
    for i, p in enumerate(d.paragraphs):
        t = (p.text or "").strip()
        if t:
            fuori.append({"id": f"p{i}", "pagina": None, "testo": t,
                          "stile": (p.style.name if p.style else ""),
                          "riquadro": None})
    for ti, tab in enumerate(d.tables):
        for ri, riga in enumerate(tab.rows):
            t = " | ".join(c.text.strip() for c in riga.cells).strip(" |")
            if t:
                fuori.append({"id": f"t{ti}r{ri}", "pagina": None, "testo": t,
                              "stile": "Tabella", "riquadro": None})
    return fuori


def _blocchi_pdf(f: Path) -> list[dict]:
    # PyMuPDF invece di pypdf: qui serve **dove** sta il testo, non solo cosa
    # dice. Senza il riquadro non si puo' evidenziare, e senza evidenziare
    # questo profilo non ha nessuna ragione di esistere.
    import fitz
    doc = fitz.open(str(f))
    fuori = []
    for n, pagina in enumerate(doc):
        for b in pagina.get_text("blocks"):
            x0, y0, x1, y1, testo = b[0], b[1], b[2], b[3], b[4]
            testo = re.sub(r"\s+", " ", (testo or "")).strip()
            if testo:
                fuori.append({"id": f"p{n}b{int(b[5])}", "pagina": n + 1,
                              "testo": testo, "stile": "",
                              "riquadro": [round(x0, 1), round(y0, 1),
                                           round(x1, 1), round(y1, 1)]})
    doc.close()
    return fuori


def _blocchi_testo(f: Path) -> list[dict]:
    righe = f.read_text(encoding="utf-8", errors="replace").splitlines()
    fuori, buffer, inizio = [], [], 0
    for i, r in enumerate(righe):
        if r.strip():
            if not buffer:
                inizio = i
            buffer.append(r.strip())
        elif buffer:
            fuori.append({"id": f"r{inizio}", "pagina": None,
                          "testo": " ".join(buffer), "stile": "",
                          "riquadro": None, "righe": i - inizio})
            buffer = []
    if buffer:
        fuori.append({"id": f"r{inizio}", "pagina": None,
                      "testo": " ".join(buffer), "stile": "",
                      "riquadro": None, "righe": len(righe) - inizio})
    return fuori


def _blocchi_righe(f: Path) -> list[dict]:
    """Il codice si taglia per righe, non per paragrafi.

    Il codice non ha righe vuote dove finisce il senso: un HTML scritto
    stretto sarebbe un blocco solo, e l'unica modifica proponibile sarebbe
    «riscrivi tutto il file» - cioe' nessuna modifica proponibile. Una riga
    per blocco e' anche l'unita' con cui si legge un errore: file, riga.
    """
    fuori = []
    for i, r in enumerate(f.read_text(encoding="utf-8",
                                      errors="replace").splitlines()):
        if r.strip():
            fuori.append({"id": f"r{i}", "pagina": None, "testo": r.rstrip(),
                          "stile": "", "riquadro": None, "righe": 1})
    return fuori


def _leggi_documento(f: Path) -> list[dict]:
    est = f.suffix.lower()
    if est == ".docx":
        return _blocchi_docx(f)
    if est == ".pdf":
        return _blocchi_pdf(f)
    if est in A_RIGHE:
        return _blocchi_righe(f)
    if est in (".txt", ".md"):
        return _blocchi_testo(f)
    raise ValueError(f"non so aprire un {est or 'file senza estensione'}. "
                     f"So aprire: {', '.join(sorted(LEGGIBILI))}")


# ------------------------------------------------------------- la sessione

def _albero(radice: Path) -> list[str]:
    """I file del progetto, in ordine, senza quelli che nessuno apre."""
    fuori = []
    for f in sorted(radice.rglob("*")):
        if len(fuori) >= ALBERO_MAX:
            break
        try:
            rel = f.relative_to(radice)
        except ValueError:
            continue
        if any(parte in NON_GUARDARE for parte in rel.parts):
            continue
        if not f.is_file():
            continue
        if f.suffix.lower() not in LEGGIBILI and f.name not in ("Makefile",):
            continue
        try:
            if f.stat().st_size > FILE_MAX:
                continue
        except OSError:
            continue
        fuori.append(str(rel).replace("\\", "/"))
    return fuori


# Da cosa si parte guardando un progetto: prima quello che si guarda, poi
# quello che si legge, poi quello che si esegue.
PRIMI = ["index.html", "README.md", "readme.md", "main.py", "app.py",
         "index.js", "main.js", "src/index.html", "src/main.py"]


def apri_cartella(percorso: str, profilo: str = "studio") -> dict:
    """Apre un progetto: l'albero a sinistra, e un file per cominciare."""
    radice = Path(percorso).expanduser().resolve()
    if not radice.is_dir():
        return {"ok": False, "motivo": f"non e' una cartella: {radice}"}
    albero = _albero(radice)
    if not albero:
        return {"ok": False,
                "motivo": f"in {radice.name} non c'e' niente che io sappia "
                          f"aprire"}
    scelto = next((x for x in PRIMI if x in albero), albero[0])
    esito = apri(str(radice / scelto), profilo=profilo, radice=str(radice),
                 albero=albero)
    if esito.get("ok"):
        esito["progetto"] = radice.name
        esito["file_nel_progetto"] = len(albero)
    return esito


def apri(percorso: str, profilo: str = "studio", radice: str = "",
         albero: list[str] | None = None) -> dict:
    """Apre un documento e lo prepara per essere cercato e indicato."""
    if profilo not in PROFILI:
        return {"ok": False,
                "motivo": f"profilo «{profilo}» sconosciuto. "
                          f"Ci sono: {', '.join(PROFILI)}"}
    f = Path(percorso).expanduser()
    if f.is_dir():
        return apri_cartella(str(f), profilo=profilo)
    if not f.is_file():
        return {"ok": False, "motivo": f"non trovo il file: {f}"}
    try:
        blocchi = _leggi_documento(f)
    except Exception as e:
        return {"ok": False, "motivo": f"{type(e).__name__}: {e}"}
    if not blocchi:
        return {"ok": False,
                "motivo": "il documento non contiene testo estraibile "
                          "(se e' una scansione servirebbe un OCR)"}

    sessione = uuid.uuid4().hex[:8]
    # Passando da un file all'altro dell'albero, il progetto non si
    # riapre: si eredita, altrimenti cliccare un file nella colonna di
    # sinistra farebbe sparire la colonna di sinistra.
    if not radice:
        vecchio = _stato() or {}
        vecchia_radice = vecchio.get("radice") or ""
        if vecchia_radice:
            try:
                f.resolve().relative_to(Path(vecchia_radice))
                radice = vecchia_radice
                albero = albero or vecchio.get("albero")
            except ValueError:
                pass
    stato = {
        "sessione": sessione,
        "profilo": profilo,
        "radice": radice,
        "albero": albero or [],
        "file": str(f.resolve()),
        "nome": f.name,
        "aperto": time.time(),
        "blocchi": blocchi,
        "evidenziati": [],
    }
    _base().mkdir(parents=True, exist_ok=True)
    (_base() / f"{sessione}.json").write_text(
        json.dumps(stato, ensure_ascii=False), encoding="utf-8")
    _puntatore().write_text(json.dumps(
        {"sessione": sessione, "profilo": profilo, "file": str(f.resolve()),
         "quando": time.time()}, ensure_ascii=False), encoding="utf-8")
    _annota(sessione, "aperto", file=str(f.resolve()), blocchi=len(blocchi),
            profilo=profilo)
    # La finestra si accende da se'. Se non ci riesce non e' un guasto del
    # lavoro - il documento e' aperto lo stesso e la ricerca funziona - ma si
    # perde la meta' che si guarda, e va detto con il motivo, non con un
    # booleano che nasconde uno schermo vuoto.
    try:
        from .harness_finestra import apri_se_serve
        finestra = apri_se_serve()
    except Exception as e:
        finestra = {"viva": False, "accesa_adesso": False,
                    "motivo": f"{type(e).__name__}: {e}"}
    pagine = {b["pagina"] for b in blocchi if b["pagina"]}
    return {"ok": True, "sessione": sessione, "nome": f.name,
            "blocchi": len(blocchi), "pagine": len(pagine) or None,
            "caratteri": sum(len(b["testo"]) for b in blocchi),
            "finestra": finestra}


def _stato(sessione: str = "") -> dict | None:
    if not sessione:
        try:
            sessione = json.loads(_puntatore().read_text(encoding="utf-8"))["sessione"]
        except Exception:
            return None
    f = _base() / f"{sessione}.json"
    try:
        return json.loads(f.read_text(encoding="utf-8"))
    except Exception:
        return None


def _salva(stato: dict) -> None:
    try:
        (_base() / f"{stato['sessione']}.json").write_text(
            json.dumps(stato, ensure_ascii=False), encoding="utf-8")
    except Exception:
        pass


def stato(sessione: str = "") -> dict:
    s = _stato(sessione)
    if s is None:
        return {"ok": False, "motivo": "nessun documento aperto nell'harness"}
    return {"ok": True, "sessione": s["sessione"], "nome": s["nome"],
            "file": s["file"], "profilo": s["profilo"],
            "blocchi": len(s["blocchi"]),
            "evidenziati": s.get("evidenziati", [])}


# ------------------------------------------------------------- la ricerca

def _parole(testo: str) -> list[str]:
    from .ricette import VUOTE
    t = unicodedata.normalize("NFKD", (testo or "").lower())
    t = "".join(c for c in t if not unicodedata.combining(c))
    return [p for p in re.findall(r"[a-z0-9]+", t)
            if len(p) > 2 and p not in VUOTE]


def _punteggio(chieste: list[str], blocco: dict) -> float:
    """Quanto un blocco risponde alla domanda.

    Si riusa la stessa uguaglianza fra parole delle ricette - contenimento,
    radice, trigrammi - cosi' la tolleranza ai refusi vale anche qui: chi
    cerca «Calhanoglu» scritto storto lo trova lo stesso.
    """
    from .ricette import _stessa_parola
    if not chieste:
        return 0.0
    dentro = set(_parole(blocco.get("testo", "")))
    if not dentro:
        return 0.0
    presi = sum(1 for p in set(chieste) if any(_stessa_parola(p, q) for q in dentro))
    return presi / len(set(chieste))


def cerca(domanda: str, quanti: int = 5, sessione: str = "") -> dict:
    """Dove sta, nel documento, quello che si sta chiedendo."""
    s = _stato(sessione)
    if s is None:
        return {"ok": False, "motivo": "nessun documento aperto nell'harness"}
    chieste = _parole(domanda)
    if not chieste:
        return {"ok": False, "motivo": "la domanda non ha parole utili"}
    punteggi = []
    for b in s["blocchi"]:
        p = _punteggio(chieste, b)
        if p > 0:
            punteggi.append((p, b))
    punteggi.sort(key=lambda x: x[0], reverse=True)
    trovati = [{"id": b["id"], "pagina": b["pagina"],
                "quanto": round(p, 2),
                "testo": b["testo"][:300]}
               for p, b in punteggi[:max(1, quanti)]]
    # Si evidenzia quello che si e' trovato: la finestra segue il registro,
    # non riceve ordini. Cosi' non c'e' niente da tenere in vita fra i due.
    s["evidenziati"] = [t["id"] for t in trovati]
    _salva(s)
    _annota(s["sessione"], "cercato", domanda=domanda[:200],
            trovati=[t["id"] for t in trovati])
    if not trovati:
        return {"ok": True, "trovati": [],
                "nota": "nel documento non c'e' niente che somigli a questo"}
    return {"ok": True, "trovati": trovati}


def leggi(intorno: str = "", blocchi: int = 3, sessione: str = "",
          caratteri: int = 4000) -> dict:
    """Il testo attorno a un punto: il contesto in cui quella cosa sta."""
    s = _stato(sessione)
    if s is None:
        return {"ok": False, "motivo": "nessun documento aperto nell'harness"}
    tutti = s["blocchi"]
    if not intorno:
        scelti = tutti[:max(1, blocchi)]
    else:
        indice = next((i for i, b in enumerate(tutti) if b["id"] == intorno), None)
        if indice is None:
            return {"ok": False, "motivo": f"nel documento non c'e' nessun «{intorno}»"}
        a = max(0, indice - blocchi)
        scelti = tutti[a:indice + blocchi + 1]
    s["evidenziati"] = [b["id"] for b in scelti]
    _salva(s)
    _annota(s["sessione"], "letto", intorno=intorno, blocchi=len(scelti))
    testo, quanti = [], 0
    for b in scelti:
        pezzo = (f"[{b['id']}"
                 + (f", pagina {b['pagina']}" if b["pagina"] else "")
                 + f"] {b['testo']}")
        if quanti + len(pezzo) > caratteri:
            break
        testo.append(pezzo)
        quanti += len(pezzo)
    return {"ok": True, "testo": "\n\n".join(testo), "blocchi": len(testo)}


# Quante pagine si disegnano al massimo. Un PDF di trecento pagine
# disegnato tutto e' mezzo giga di immagini per niente: chi legge ne guarda
# poche, e quelle che servono le indica la ricerca.
PAGINE_MAX = 40
ZOOM = 1.7


def pagine_disegnate(sessione: str = "") -> list[dict]:
    """Disegna le pagine del PDF, con sopra le evidenziazioni.

    Torna [{pagina, file, evidenziata}]. Per i formati che non sono PDF
    torna [] - li' non c'e' una pagina da disegnare, e fingerla sarebbe
    peggio del testo.
    """
    s = _stato(sessione)
    if s is None or not s["file"].lower().endswith(".pdf"):
        return []
    import fitz
    cartella = _base() / f"{s['sessione']}-pagine"
    cartella.mkdir(parents=True, exist_ok=True)
    accesi = set(s.get("evidenziati") or [])
    per_pagina: dict[int, list[list[float]]] = {}
    for b in s["blocchi"]:
        if b["id"] in accesi and b.get("riquadro") and b.get("pagina"):
            per_pagina.setdefault(b["pagina"], []).append(b["riquadro"])

    fuori = []
    try:
        doc = fitz.open(s["file"])
    except Exception:
        return []
    try:
        for n in range(min(len(doc), PAGINE_MAX)):
            numero = n + 1
            riquadri = per_pagina.get(numero, [])
            # Il nome porta dentro le evidenziazioni: cosi' una pagina
            # gia' disegnata con le stesse non si ridisegna, e una con
            # evidenziazioni diverse non viene scambiata per quella.
            firma = "-".join(f"{int(r[0])}_{int(r[1])}" for r in riquadri) or "pulita"
            f = cartella / f"p{numero}-{firma}.png"
            if not f.exists():
                pagina = doc[n]
                for r in riquadri:
                    # Un'annotazione di evidenziazione vera, disegnata da chi
                    # possiede le coordinate. Nessuna conversione fra mondi.
                    ann = pagina.add_highlight_annot(fitz.Rect(*r))
                    ann.set_colors(stroke=(0.91, 0.45, 0.29))   # la brace
                    ann.set_opacity(0.28)
                    ann.update()
                pagina.get_pixmap(matrix=fitz.Matrix(ZOOM, ZOOM)).save(str(f))
            fuori.append({"pagina": numero, "file": str(f),
                          "evidenziata": bool(riquadri)})
    finally:
        doc.close()
    return fuori


def _indice_file(radice: Path) -> Path:
    impronta = hashlib.sha1(
        str(radice).lower().encode("utf-8")).hexdigest()[:12]
    return _base() / f"indice-{impronta}.json"


def _indice(radice: Path, albero: list[str]) -> dict:
    """I blocchi di tutti i file del progetto, tenuti da parte.

    Rileggere sei PDF a ogni domanda costerebbe secondi ogni volta, e chi
    studia fa molte domande di fila. L'indice si aggiorna per file, guardando
    la data: un libro che non e' cambiato non si riapre.
    """
    f = _indice_file(radice)
    try:
        vecchio = json.loads(f.read_text(encoding="utf-8"))
    except Exception:                                          # noqa: BLE001
        vecchio = {}
    nuovo, cambiato = {}, False
    for rel in albero:
        p = radice / rel
        try:
            quando = p.stat().st_mtime
        except OSError:
            continue
        prima = vecchio.get(rel)
        if prima and abs(prima.get("quando", 0) - quando) < 0.001:
            nuovo[rel] = prima
            continue
        try:
            nuovo[rel] = {"quando": quando,
                          "blocchi": _leggi_documento(p)}
            cambiato = True
        except Exception:                                      # noqa: BLE001
            # Un file che non si sa leggere non ferma la ricerca negli altri.
            nuovo[rel] = {"quando": quando, "blocchi": []}
            cambiato = True
    if cambiato or set(nuovo) != set(vecchio):
        try:
            f.parent.mkdir(parents=True, exist_ok=True)
            f.write_text(json.dumps(nuovo, ensure_ascii=False),
                         encoding="utf-8")
        except Exception:                                      # noqa: BLE001
            pass
    return nuovo


def cerca_progetto(domanda: str, quanti: int = 8,
                   sessione: str = "") -> dict:
    """Cerca in tutti i file del progetto, non solo in quello aperto.

    E' la differenza fra leggere un libro e studiare su sei: la domanda
    «dove si parla di X» non ha senso dentro un documento solo, quando i
    documenti sono una pila. La risposta dice il file **e** la pagina,
    perche' una citazione senza il posto non si puo' controllare.
    """
    stato = _stato(sessione)
    if not stato:
        return {"ok": False, "motivo": "non c'e' niente di aperto"}
    radice = stato.get("radice") or ""
    albero = stato.get("albero") or []
    if not radice or not albero:
        return {"ok": False,
                "motivo": "non c'e' un progetto aperto: apri una cartella "
                          "invece di un file solo"}
    chieste = _parole(domanda)
    if not chieste:
        return {"ok": False, "motivo": "la domanda e' vuota"}
    indice = _indice(Path(radice), albero)
    trovati = []
    for rel, dati in indice.items():
        for b in dati.get("blocchi", []):
            punti = _punteggio(chieste, b)
            if punti > 0:
                trovati.append((punti, rel, b))
    trovati.sort(key=lambda x: -x[0])
    dentro = [{"file": rel, "blocco": b["id"], "pagina": b.get("pagina"),
               "testo": b["testo"][:400], "punti": round(p, 3)}
              for p, rel, b in trovati[:quanti]]
    _annota(stato["sessione"], "cerca nel progetto", domanda=domanda,
            trovati=len(dentro))
    return {"ok": True, "domanda": domanda, "quanti": len(dentro),
            "cercati": len(indice), "risultati": dentro}


def eventi(sessione: str = "", quanti: int = 50) -> list[dict]:
    """Cosa e' stato guardato, in ordine. Il racconto non serve: c'e' il registro."""
    s = _stato(sessione)
    if s is None:
        return []
    f = _registro(s["sessione"])
    if not f.exists():
        return []
    fuori = []
    for riga in f.read_text(encoding="utf-8").splitlines():
        riga = riga.strip()
        if riga:
            try:
                fuori.append(json.loads(riga))
            except Exception:
                pass
    return fuori[-quanti:]


def chiudi(sessione: str = "") -> bool:
    s = _stato(sessione)
    if s is None:
        return False
    _annota(s["sessione"], "chiuso")
    try:
        _puntatore().unlink()
    except Exception:
        pass
    return True
