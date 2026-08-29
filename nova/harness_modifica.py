# -*- coding: utf-8 -*-
"""NOVA scrive dentro il documento, ma non di nascosto.

Il pezzo che mancava all'harness era questo: fin qui NOVA sapeva leggere e
indicare, non toccare. Toccare pero' e' l'azione che non si annulla da se',
quindi qui non esiste una funzione che modifichi e basta. Esiste una
proposta, che sta in un file accanto alla sessione e non tocca niente, e
esiste un'applicazione, che l'utente chiede dopo aver visto cosa cambia.

Le due meta' sono separate apposta: una proposta si puo' guardare, discutere
e buttare senza conseguenze, e chi applica sa sempre cosa sta applicando.

Sui formati non si promette quello che non si sa mantenere:

  .md .txt   si riscrivono per intero, e' testo, non si perde niente.
  .html      come sopra, ma tagliato per righe invece che per paragrafi: il
             codice non ha righe vuote dove finisce il senso. La finestra lo
             mostra reso - e' un artifact, si guarda per quello che fa - e il
             sorgente e' a un click, perche' e' l'unica cosa che si salva.
  .docx      si modifica il paragrafo, non il documento: python-docx
             riscrive il testo dentro la prima porzione e cancella le
             altre, cosi' il carattere, lo stile e il resto della pagina
             restano quelli che erano. Rifare un .docx da zero a partire dal
             testo estratto sarebbe stato molto piu' facile e avrebbe buttato
             via l'impaginazione di chi lo ha scritto.
  .pdf       il testo non si riscrive. Un PDF non contiene paragrafi, contiene
             lettere messe in un punto della pagina: cambiarne una vuol dire
             ridisegnare quello che c'e' intorno, e il risultato si vede.
             Si annota pero' per davvero - evidenziazioni e note gialle che
             restano nel file e si aprono in qualunque lettore.
"""
from __future__ import annotations

import hashlib
import json
import shutil
import time
from pathlib import Path

from . import harness
from . import registro

AZIONI_TESTO = {"sostituisci", "prima", "dopo", "elimina"}
AZIONI_PDF = {"evidenzia", "nota"}
# Togliere e evidenziare sono gesti, non testi: chiederglielo sarebbe una
# domanda senza risposta possibile.
SENZA_TESTO = {"elimina", "evidenzia"}
ESTRATTO = 220


def file_proposta(percorso: str) -> Path:
    """La proposta si lega al documento, non alla sessione.

    Legarla alla sessione sembrava naturale ed era sbagliato: chiudere e
    riaprire lo stesso documento crea una sessione nuova, e la proposta in
    attesa spariva senza che nessuno lo dicesse. Una modifica proposta e poi
    persa in silenzio e' peggio di una modifica rifiutata.
    """
    impronta = hashlib.sha1(
        str(Path(percorso).resolve()).lower().encode("utf-8")).hexdigest()[:12]
    return harness._base() / f"proposta-{impronta}.json"


def _corta(t: str, quanto: int = ESTRATTO) -> str:
    t = " ".join((t or "").split())
    return t if len(t) <= quanto else t[:quanto - 1] + "…"


# ------------------------------------------------------------- la proposta

def proponi(modifiche: list[dict], sessione: str = "",
            motivo: str = "") -> dict:
    """Prepara le modifiche e le mostra. Non scrive niente nel documento."""
    stato = harness._stato(sessione)
    if not stato:
        return {"ok": False, "motivo": "non c'e' nessun documento aperto"}
    f = Path(stato["file"])
    if not f.is_file():
        return {"ok": False, "motivo": f"il file non c'e' piu': {f}"}
    est = f.suffix.lower()
    lecite = AZIONI_PDF if est == ".pdf" else AZIONI_TESTO
    if not modifiche:
        return {"ok": False, "motivo": "nessuna modifica da proporre"}

    per_id = {b["id"]: b for b in stato["blocchi"]}
    pronte, guai = [], []
    for n, m in enumerate(modifiche):
        azione = (m.get("azione") or "sostituisci").strip().lower()
        blocco = (m.get("blocco") or "").strip()
        testo = m.get("testo") or ""
        if azione not in lecite:
            guai.append(f"modifica {n + 1}: su un {est} si puo' fare "
                        f"{', '.join(sorted(lecite))}, non «{azione}»")
            continue
        if blocco not in per_id:
            guai.append(f"modifica {n + 1}: il blocco «{blocco}» non "
                        f"esiste in questo documento")
            continue
        if azione not in SENZA_TESTO and not testo.strip():
            guai.append(f"modifica {n + 1}: manca il testo")
            continue
        if est == ".docx" and azione in ("prima", "dopo") \
                and blocco.startswith("t"):
            guai.append(f"modifica {n + 1}: dentro una tabella si sostituisce "
                        f"la riga, non se ne aggiungono")
            continue
        pronte.append({"azione": azione, "blocco": blocco, "testo": testo,
                       "prima": per_id[blocco]["testo"],
                       "righe": per_id[blocco].get("righe"),
                       "pagina": per_id[blocco].get("pagina")})
    if guai:
        return {"ok": False, "motivo": "; ".join(guai)}

    proposta = {
        "sessione": stato["sessione"],
        "file": str(f),
        "nome": f.name,
        "motivo": motivo,
        "quando": time.time(),
        "modifiche": pronte,
    }
    file_proposta(str(f)).write_text(
        json.dumps(proposta, ensure_ascii=False), encoding="utf-8")
    harness._annota(stato["sessione"], "proposta", quante=len(pronte),
                    motivo=motivo)
    return {"ok": True, "sessione": stato["sessione"], "quante": len(pronte),
            "in_attesa": True,
            "anteprima": [
                {"blocco": p["blocco"], "azione": p["azione"],
                 "prima": _corta(p["prima"]),
                 "dopo": (_corta(p["testo"])
                          if p["azione"] not in SENZA_TESTO else "")}
                for p in pronte],
            "nota": "nessuna riga e' ancora cambiata: serve harness_applica"}


def proposta(sessione: str = "") -> dict | None:
    stato = harness._stato(sessione)
    if not stato:
        return None
    try:
        return json.loads(
            file_proposta(stato["file"]).read_text(encoding="utf-8"))
    except Exception:                                          # noqa: BLE001
        return None


def scarta(sessione: str = "") -> dict:
    stato = harness._stato(sessione)
    if not stato:
        return {"ok": False, "motivo": "non c'e' nessun documento aperto"}
    f = file_proposta(stato["file"])
    c_era = f.exists()
    f.unlink(missing_ok=True)
    if c_era:
        harness._annota(stato["sessione"], "proposta scartata")
    return {"ok": True, "scartata": c_era}


# ---------------------------------------------------------- l'applicazione

def applica(sessione: str = "") -> dict:
    """Scrive davvero. Prima mette da parte una copia intatta."""
    p = proposta(sessione)
    if not p:
        return {"ok": False, "motivo": "non c'e' nessuna proposta da applicare"}
    f = Path(p["file"])
    if not f.is_file():
        return {"ok": False, "motivo": f"il file non c'e' piu': {f}"}

    # La copia viene prima della scrittura, non dopo: se la scrittura fallisce
    # a meta' la copia c'e' gia', ed e' esattamente il caso in cui serve.
    copia = f.with_suffix(f.suffix + ".prima")
    try:
        shutil.copy2(f, copia)
    except Exception as e:                                     # noqa: BLE001
        return {"ok": False,
                "motivo": f"non riesco a mettere da parte una copia "
                          f"({type(e).__name__}: {e}), quindi non tocco niente"}

    est = f.suffix.lower()
    try:
        if est in (".md", ".txt", ".html", ".htm"):
            fatte = _applica_testo(f, p["modifiche"])
        elif est == ".docx":
            fatte = _applica_docx(f, p["modifiche"])
        elif est == ".pdf":
            fatte = _applica_pdf(f, p["modifiche"])
        else:
            return {"ok": False, "motivo": f"non so scrivere dentro un {est}"}
    except Exception as e:                                     # noqa: BLE001
        try:
            shutil.copy2(copia, f)
        except Exception:                                      # noqa: BLE001
            pass
        return {"ok": False,
                "motivo": f"{type(e).__name__}: {e} — il documento e' "
                          f"stato rimesso com'era"}

    file_proposta(p["file"]).unlink(missing_ok=True)
    _rileggi(p["sessione"], f)
    harness._annota(p["sessione"], "applicata", quante=fatte, file=str(f))
    registro.annota("modificato un documento", dove=str(f),
                    dettagli=f"{fatte} modifiche; copia intatta in "
                             f"{copia.name}",
                    tipo="documento", esito="ok")
    return {"ok": True, "applicate": fatte, "file": str(f),
            "copia_di_prima": str(copia)}


def _rileggi(sessione: str, f: Path) -> None:
    """Il documento e' cambiato: i blocchi di prima non valgono piu'."""
    stato = harness._stato(sessione)
    if not stato:
        return
    try:
        stato["blocchi"] = harness._leggi_documento(f)
    except Exception:                                          # noqa: BLE001
        return
    validi = {b["id"] for b in stato["blocchi"]}
    stato["evidenziati"] = [x for x in stato.get("evidenziati", [])
                            if x in validi]
    harness._salva(stato)


def _rifai(righe: list[str], modifiche: list[dict],
           nuovo: str = "", vecchio: str = "") -> tuple[list[str], int]:
    """Le righe come saranno. Con le marche, come si vedono prima.

    Lo stesso codice serve l'anteprima e l'applicazione, e non per pigrizia:
    un'anteprima calcolata a parte prima o poi mostra qualcosa di diverso da
    quello che poi succede, ed e' il modo piu' sicuro di far perdere fiducia
    a chi deve premere il bottone. Con le marche in coda alle righe - una per
    cio' che arriva, una per cio' che se ne va - il testo e' lo stesso, e chi
    lo disegna sa cosa colorare.
    """
    # Si lavora dal fondo verso l'alto: cosi' gli indici delle modifiche
    # ancora da fare restano quelli calcolati all'apertura.
    ordinate = sorted(modifiche, key=lambda m: _inizio(m["blocco"]) or 0,
                      reverse=True)
    anteprima = bool(nuovo or vecchio)
    fatte = 0
    for m in ordinate:
        i = _inizio(m["blocco"])
        if i is None or i >= len(righe):
            continue
        # Quante righe occupa il blocco lo sa il blocco: nei documenti a
        # paragrafi si arriva fino alla riga vuota; nel codice il blocco e'
        # una riga sola, e fermarsi alla riga vuota si mangerebbe mezzo file.
        quante = m.get("righe")
        if quante:
            fine = min(len(righe), i + quante)
        else:
            fine = i
            while fine < len(righe) and righe[fine].strip():
                fine += 1
        nuove = [r + nuovo for r in (m["testo"] or "").splitlines()] \
            if m["testo"] else []
        vecchie = [righe[k] + vecchio for k in range(i, fine)]
        if m["azione"] == "sostituisci":
            righe[i:fine] = (vecchie + [""] + nuove) if anteprima else nuove
        elif m["azione"] == "elimina":
            # La riga vuota dopo il blocco se ne va con lui, ma solo se e'
            # davvero vuota: nel codice, la riga dopo e' altro codice.
            coda = (fine + 1 if fine < len(righe) and not righe[fine].strip()
                    else fine)
            righe[i:coda] = (vecchie + [""]) if anteprima else []
        elif m["azione"] == "prima":
            righe[i:i] = nuove + [""]
        elif m["azione"] == "dopo":
            righe[fine:fine] = [""] + nuove
        fatte += 1
    return righe, fatte


def anteprima_testo(f: Path, modifiche: list[dict],
                    nuovo: str, vecchio: str) -> str:
    """Il documento come sarebbe, con le marche su cio' che cambia."""
    righe, _ = _rifai(f.read_text(encoding="utf-8").splitlines(),
                      modifiche, nuovo, vecchio)
    return "\n".join(righe)


def _applica_testo(f: Path, modifiche: list[dict]) -> int:
    righe, fatte = _rifai(f.read_text(encoding="utf-8").splitlines(),
                          modifiche)
    testo = "\n".join(righe)
    if not testo.endswith("\n"):
        testo += "\n"
    f.write_text(testo, encoding="utf-8")
    return fatte


def _inizio(blocco: str) -> int | None:
    if not blocco.startswith("r"):
        return None
    try:
        return int(blocco[1:])
    except ValueError:
        return None


def _applica_docx(f: Path, modifiche: list[dict]) -> int:
    import docx
    d = docx.Document(str(f))
    fatte = 0
    for m in modifiche:
        b = m["blocco"]
        if b.startswith("t"):
            if _riga_tabella(d, b, m):
                fatte += 1
            continue
        try:
            i = int(b[1:])
            par = d.paragraphs[i]
        except (ValueError, IndexError):
            continue
        if m["azione"] == "sostituisci":
            _scrivi_paragrafo(par, m["testo"])
        elif m["azione"] == "elimina":
            par._element.getparent().remove(par._element)
        elif m["azione"] == "prima":
            nuovo = par.insert_paragraph_before(m["testo"])
            nuovo.style = par.style
        elif m["azione"] == "dopo":
            dopo = d.paragraphs[i + 1] if i + 1 < len(d.paragraphs) else None
            if dopo is not None:
                nuovo = dopo.insert_paragraph_before(m["testo"])
                nuovo.style = par.style
            else:
                d.add_paragraph(m["testo"], style=par.style)
        fatte += 1
    d.save(str(f))
    return fatte


def _scrivi_paragrafo(par, testo: str) -> None:
    """Il testo cambia, il vestito no.

    La prima porzione tiene il carattere, il corpo, il colore: ci si scrive
    dentro e si tolgono le altre. Cancellare il paragrafo e riscriverlo
    avrebbe perso tutto quello che non e' testo.
    """
    if par.runs:
        par.runs[0].text = testo
        for r in par.runs[1:]:
            r.text = ""
    else:
        par.add_run(testo)


def _riga_tabella(d, blocco: str, m: dict) -> bool:
    if m["azione"] != "sostituisci":
        return False
    try:
        ti, ri = blocco[1:].split("r")
        riga = d.tables[int(ti)].rows[int(ri)]
    except Exception:                                          # noqa: BLE001
        return False
    celle = [c.strip() for c in m["testo"].split("|")]
    for cella, nuovo in zip(riga.cells, celle):
        par = cella.paragraphs[0] if cella.paragraphs else cella.add_paragraph()
        _scrivi_paragrafo(par, nuovo)
        for extra in cella.paragraphs[1:]:
            extra._element.getparent().remove(extra._element)
    return True


def _applica_pdf(f: Path, modifiche: list[dict]) -> int:
    import fitz
    doc = fitz.open(str(f))
    fatte = 0
    for m in modifiche:
        b = m["blocco"]
        if "b" not in b or not b.startswith("p"):
            continue
        try:
            n = int(b[1:b.index("b")])
            pagina = doc[n]
        except Exception:                                      # noqa: BLE001
            continue
        riquadro = _riquadro(doc, b)
        if riquadro is None:
            continue
        if m["azione"] == "evidenzia":
            pagina.add_highlight_annot(riquadro)
        elif m["azione"] == "nota":
            nota = pagina.add_text_annot(
                fitz.Point(riquadro.x1 + 4, riquadro.y0), m["testo"])
            nota.set_info(title="NOVA")
            nota.update()
        fatte += 1
    doc.saveIncr() if doc.can_save_incrementally() else doc.save(
        str(f), incremental=False)
    doc.close()
    return fatte


def _riquadro(doc, blocco: str):
    import fitz
    try:
        n = int(blocco[1:blocco.index("b")])
        num = int(blocco[blocco.index("b") + 1:])
    except Exception:                                          # noqa: BLE001
        return None
    for b in doc[n].get_text("blocks"):
        if int(b[5]) == num:
            return fitz.Rect(b[0], b[1], b[2], b[3])
    return None
