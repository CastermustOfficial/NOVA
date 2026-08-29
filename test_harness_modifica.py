# -*- coding: utf-8 -*-
"""Scrivere dentro un documento e' l'azione che non si annulla da se'.

Quindi qui si prova soprattutto quello che deve NON succedere: che una
proposta cambi qualcosa prima che qualcuno l'abbia vista, che applicare
lasci il documento senza una copia intatta accanto, che un errore a meta'
strada lasci il file mezzo riscritto, e che un .docx torni indietro come
testo semplice avendo perso il grassetto di chi l'aveva scritto.

Il .pdf ha una regola sua e va provata come tale: il testo non si riscrive,
si annota. Promettere una riscrittura e consegnare un pasticcio sarebbe
stato peggio che dire di no.
"""
import os
import sys
import tempfile
from pathlib import Path

RADICE = Path(__file__).resolve().parent
sys.path.insert(0, str(RADICE))
sys.stdout.reconfigure(encoding="utf-8", errors="replace")

finto = Path(tempfile.mkdtemp(prefix="nova_mod_"))
os.environ["APPDATA"] = str(finto)

passati = 0
falliti: list[str] = []


def controlla(nome, condizione, dettaglio=""):
    global passati
    if condizione:
        passati += 1
        print(f"  [ok ] {nome}")
    else:
        falliti.append(nome)
        print(f"  [NO ] {nome}  {dettaglio}")


from nova import harness                      # noqa: E402
from nova import harness_modifica as mod      # noqa: E402

# La finestra non si apre: qui si prova il testo, non il vetro.
import nova.harness_finestra as hf            # noqa: E402
hf.apri_se_serve = lambda *a, **k: {"viva": False, "accesa_adesso": False}

lavoro = Path(tempfile.mkdtemp(prefix="nova_doc_"))

print("\n1. una proposta non tocca niente")
md = lavoro / "note.md"
originale = ("# Relazione\n\nIl primo paragrafo dice una cosa sbagliata.\n\n"
             "Il secondo invece va bene.\n")
md.write_text(originale, encoding="utf-8")
harness.apri(str(md))
blocchi = harness._stato()["blocchi"]
sbagliato = [b for b in blocchi if "sbagliata" in b["testo"]][0]

r = mod.proponi([{"blocco": sbagliato["id"], "azione": "sostituisci",
                  "testo": "Il primo paragrafo ora dice la cosa giusta."}],
                motivo="correzione")
controlla("la proposta si prepara", r.get("ok"), str(r.get("motivo")))
controlla("e dice che nessuno ha ancora scritto", r.get("in_attesa") is True)
controlla("il file sul disco e' ancora quello di prima",
          md.read_text(encoding="utf-8") == originale)
controlla("c'e' il prima e il dopo, non solo il dopo",
          "sbagliata" in r["anteprima"][0]["prima"]
          and "giusta" in r["anteprima"][0]["dopo"], str(r["anteprima"]))
controlla("e la proposta si ritrova da fuori",
          (mod.proposta() or {}).get("motivo") == "correzione")

print("\n2. un blocco che non esiste si rifiuta prima, non a meta' strada")
brutta = mod.proponi([{"blocco": "r999", "azione": "sostituisci", "testo": "x"}])
controlla("rifiutata", brutta.get("ok") is False)
controlla("con il motivo, non con un falso", "non esiste" in brutta["motivo"],
          brutta.get("motivo", ""))
controlla("e la proposta buona di prima e' ancora li'",
          (mod.proposta() or {}).get("motivo") == "correzione")

print("\n3. applicare scrive, e lascia accanto com'era")
a = mod.applica()
testo = md.read_text(encoding="utf-8")
controlla("applicata", a.get("ok"), str(a.get("motivo")))
controlla("la riga nuova c'e'", "la cosa giusta" in testo, testo)
controlla("quella vecchia non c'e' piu'", "sbagliata" not in testo, testo)
controlla("il titolo non e' stato toccato", testo.startswith("# Relazione"),
          testo[:30])
controlla("e nemmeno il secondo paragrafo", "Il secondo invece va bene." in testo)
controlla("la copia di prima esiste",
          md.with_suffix(".md.prima").exists())
controlla("e contiene davvero l'originale",
          md.with_suffix(".md.prima").read_text(encoding="utf-8") == originale)
controlla("la proposta e' sparita dopo l'uso", mod.proposta() is None)
controlla("i blocchi sono stati riletti",
          any("giusta" in b["testo"] for b in harness._stato()["blocchi"]))

print("\n4. aggiungere e togliere")
b2 = harness._stato()["blocchi"]
ultimo = b2[-1]
mod.proponi([{"blocco": ultimo["id"], "azione": "dopo",
              "testo": "Un paragrafo aggiunto in coda."}])
mod.applica()
testo = md.read_text(encoding="utf-8")
controlla("l'aggiunta e' in fondo", testo.rstrip().endswith("in coda."), testo)
controlla("e ha una riga vuota che la separa",
          "va bene.\n\nUn paragrafo aggiunto" in testo, repr(testo))

b3 = harness._stato()["blocchi"]
primo_par = [b for b in b3 if "giusta" in b["testo"]][0]
mod.proponi([{"blocco": primo_par["id"], "azione": "elimina"}])
mod.applica()
testo = md.read_text(encoding="utf-8")
controlla("eliminare toglie il paragrafo", "la cosa giusta" not in testo, testo)
controlla("senza lasciare un buco doppio", "\n\n\n" not in testo, repr(testo))

print("\n5. scartare non lascia strascichi")
prima = md.read_text(encoding="utf-8")
mod.proponi([{"blocco": harness._stato()["blocchi"][0]["id"],
              "azione": "sostituisci", "testo": "MAI"}])
s = mod.scarta()
controlla("scartata", s.get("scartata") is True)
controlla("il file non e' cambiato", md.read_text(encoding="utf-8") == prima)
controlla("e non c'e' piu' niente in attesa", mod.proposta() is None)
controlla("applicare adesso non fa danni",
          mod.applica().get("ok") is False)

print("\n6. un .docx si modifica senza spogliarlo")
try:
    import docx
    from docx.shared import Pt
    d = docx.Document()
    d.add_heading("Titolo del documento", level=1)
    p = d.add_paragraph()
    run = p.add_run("Questa frase e' in grassetto e va corretta.")
    run.bold = True
    run.font.size = Pt(16)
    d.add_paragraph("Una frase normale che resta dov'e'.")
    tab = d.add_table(rows=2, cols=2)
    tab.cell(0, 0).text = "Voce"
    tab.cell(0, 1).text = "Valore"
    tab.cell(1, 0).text = "Alfa"
    tab.cell(1, 1).text = "uno"
    doc = lavoro / "relazione.docx"
    d.save(str(doc))

    harness.apri(str(doc))
    bd = harness._stato()["blocchi"]
    da_correggere = [b for b in bd if "va corretta" in b["testo"]][0]
    mod.proponi([{"blocco": da_correggere["id"], "azione": "sostituisci",
                  "testo": "Questa frase e' in grassetto ed e' corretta."}])
    e = mod.applica()
    controlla("il docx si applica", e.get("ok"), str(e.get("motivo")))

    d2 = docx.Document(str(doc))
    par = [x for x in d2.paragraphs if "corretta" in x.text][0]
    controlla("il testo e' cambiato", "ed e' corretta" in par.text, par.text)
    controlla("il grassetto e' sopravvissuto", par.runs[0].bold is True)
    controlla("e anche il corpo del carattere",
              par.runs[0].font.size == Pt(16), str(par.runs[0].font.size))
    controlla("il titolo e' ancora un titolo",
              d2.paragraphs[0].style.name.startswith("Heading"),
              d2.paragraphs[0].style.name)
    controlla("l'altro paragrafo non e' stato toccato",
              any("resta dov'e'" in x.text for x in d2.paragraphs))
    controlla("la tabella e' ancora una tabella", len(d2.tables) == 1)
    controlla("la copia di prima c'e'",
              doc.with_suffix(".docx.prima").exists())

    b4 = harness._stato()["blocchi"]
    riga = [b for b in b4 if b["id"].startswith("t") and "Alfa" in b["testo"]][0]
    mod.proponi([{"blocco": riga["id"], "azione": "sostituisci",
                  "testo": "Beta | due"}])
    mod.applica()
    d3 = docx.Document(str(doc))
    controlla("una riga di tabella si riscrive",
              d3.tables[0].cell(1, 0).text == "Beta"
              and d3.tables[0].cell(1, 1).text == "due",
              d3.tables[0].cell(1, 0).text)
    controlla("l'intestazione della tabella e' intatta",
              d3.tables[0].cell(0, 0).text == "Voce")

    dentro = mod.proponi([{"blocco": riga["id"], "azione": "dopo",
                           "testo": "x"}])
    controlla("dentro una tabella non si aggiungono righe alla cieca",
              dentro.get("ok") is False, str(dentro))
except ImportError:
    print("  (python-docx non c'e': salto)")

print("\n7. un .pdf si annota, non si riscrive")
try:
    import fitz
    pdf = lavoro / "articolo.pdf"
    doc = fitz.open()
    pagina = doc.new_page()
    pagina.insert_text((72, 100), "Lo stato dell'arte dei modelli aperti.")
    pagina.insert_text((72, 130), "La quantizzazione riduce la memoria.")
    doc.save(str(pdf))
    doc.close()

    harness.apri(str(pdf))
    bp = harness._stato()["blocchi"]
    quanti_prima = len(bp)
    bersaglio = [b for b in bp if "quantizzazione" in b["testo"]][0]

    no = mod.proponi([{"blocco": bersaglio["id"], "azione": "sostituisci",
                       "testo": "non si puo'"}])
    controlla("riscrivere il testo di un pdf si rifiuta", no.get("ok") is False)
    controlla("e si dice cosa invece si puo' fare",
              "evidenzia" in no.get("motivo", ""), no.get("motivo", ""))

    si = mod.proponi([{"blocco": bersaglio["id"], "azione": "evidenzia"},
                      {"blocco": bp[0]["id"], "azione": "nota",
                       "testo": "da citare"}])
    controlla("evidenziare e annotare si accettano", si.get("ok"), str(si))
    e = mod.applica()
    controlla("e si applicano davvero", e.get("ok"), str(e.get("motivo")))

    riletto = fitz.open(str(pdf))
    tipi = [a.type[1] for p in riletto for a in p.annots()]
    riletto.close()
    controlla("nel pdf c'e' un'evidenziazione", "Highlight" in tipi, str(tipi))
    controlla("e una nota", "Text" in tipi, str(tipi))
    controlla("la copia di prima c'e'", pdf.with_suffix(".pdf.prima").exists())
    controlla("il testo del pdf e' rimasto quello",
              len(harness._stato()["blocchi"]) == quanti_prima)
except ImportError:
    print("  (PyMuPDF non c'e': salto)")

print("\n8. senza niente aperto non si propone")
harness.chiudi()
v = mod.proponi([{"blocco": "r0", "azione": "sostituisci", "testo": "x"}])
controlla("rifiutata", v.get("ok") is False)
controlla("con il motivo giusto", "nessun documento" in v.get("motivo", ""),
          v.get("motivo", ""))

print(f"\n{passati}/{passati + len(falliti)} passati")
for x in falliti:
    print("  FALLITO:", x)
sys.exit(1 if falliti else 0)
